"""
Run this script once to generate Project_Documentation.docx
    python generate_docs.py
Requires:  pip install python-docx
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.8)
    section.right_margin  = Cm(2.8)

# ── Styles helpers ────────────────────────────────────────────────────────────
def set_font(run, name="Calibri", size=11, bold=False, italic=False, color=None):
    run.font.name  = name
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.name = "Calibri"
        if level == 1:
            run.font.color.rgb = RGBColor(31, 73, 125)
        elif level == 2:
            run.font.color.rgb = RGBColor(68, 114, 196)
        elif level == 3:
            run.font.color.rgb = RGBColor(31, 73, 125)
    return h

def para(text, bold=False, size=11, color=None, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.3 + level * 0.3)
    run = p.add_run(text)
    set_font(run, size=10.5)

def code_block(lines):
    """Shaded paragraph simulating a code block."""
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Inches(0.4)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        run = p.add_run(line)
        run.font.name   = "Courier New"
        run.font.size   = Pt(9.5)
        run.font.color.rgb = RGBColor(30, 30, 30)
        shading = OxmlElement("w:shd")
        shading.set(qn("w:val"), "clear")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:fill"), "F2F2F2")
        p._p.pPr.append(shading)

def table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.font.bold = True
            run.font.name = "Calibri"
            run.font.size = Pt(10)
        # header shading
        tc = hdr[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  "1F497D")
        tcPr.append(shd)
        for par in hdr[i].paragraphs:
            for run in par.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)

    for row_data in rows:
        row = t.add_row().cells
        for i, cell_text in enumerate(row_data):
            row[i].text = cell_text
            for run in row[i].paragraphs[0].runs:
                run.font.name = "Calibri"
                run.font.size = Pt(10)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    return t

def hr():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pb  = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),  "single")
    bottom.set(qn("w:sz"),   "6")
    bottom.set(qn("w:space"),"1")
    bottom.set(qn("w:color"),"4472C4")
    pb.append(bottom)
    pPr.append(pb)


# ═════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ═════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title_p.add_run("🏦  Agentic AI Loan & Credit Advisor")
r.font.name  = "Calibri"
r.font.size  = Pt(28)
r.font.bold  = True
r.font.color.rgb = RGBColor(31, 73, 125)

doc.add_paragraph()
sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub_p.add_run("Comprehensive Technical Project Documentation")
r2.font.name  = "Calibri"
r2.font.size  = Pt(16)
r2.font.color.rgb = RGBColor(68, 114, 196)

doc.add_paragraph()
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = date_p.add_run(f"Version 1.0  ·  {datetime.date.today().strftime('%B %Y')}")
r3.font.name  = "Calibri"
r3.font.size  = Pt(12)
r3.font.color.rgb = RGBColor(120, 120, 120)

doc.add_page_break()


# ═════════════════════════════════════════════════════════════════════════════
# 1. PROJECT OVERVIEW
# ═════════════════════════════════════════════════════════════════════════════
heading("1. Project Overview")
para(
    "The Agentic AI Loan & Credit Advisor is a production-grade, end-to-end intelligent financial advisory "
    "platform built for the Indian market. The system combines multiple specialized AI agents, real-time "
    "market data scraping, XGBoost machine learning, Retrieval-Augmented Generation (RAG), and a live "
    "WebSocket chat interface to deliver personalized loan recommendations in seconds."
)
doc.add_paragraph()
para(
    "Unlike simple chatbot wrappers, this system is a deterministic AI State Machine. The LLM is strictly "
    "used for language understanding and natural language generation. All financial calculations — EMI, "
    "DTI, risk scores, and interest rates — are computed by deterministic Python code, ensuring 100% "
    "mathematical accuracy with zero AI hallucinations on financial figures.",
    bold=False
)

doc.add_paragraph()
heading("1.1  Core Capabilities", level=2)
bullets = [
    "Personalized loan recommendations across 15+ loan types (Home, Personal, Car, Business, Gold, Education, MSME, MUDRA, and more)",
    "Real-time interest rate intelligence: rates are scraped from the internet at every server boot and locked in RAM",
    "XGBoost-powered credit risk scoring using 10 financial features",
    "Regulatory compliance validation against RBI guidelines, PMAY, and MUDRA schemes via RAG",
    "Multi-user, persistent chat sessions with full conversation history",
    "Document upload and automated financial profile extraction",
    "Guest/demo mode for unauthenticated users",
    "Glassmorphism dark-themed responsive React frontend with real-time WebSocket chat",
]
for b in bullets:
    bullet(b)


# ═════════════════════════════════════════════════════════════════════════════
# 2. TECH STACK
# ═════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading("2. Technology Stack")

heading("2.1  Backend", level=2)
table(
    ["Layer", "Technology", "Purpose"],
    [
        ["Web Framework",        "FastAPI 0.135 + Uvicorn",                 "Async HTTP + WebSocket server"],
        ["AI Orchestration",     "LangGraph 1.0",                           "Stateful multi-node AI workflow graph"],
        ["Agentic Framework",    "CrewAI 1.10",                             "Specialized role-based AI agents with tool calling"],
        ["LLM Provider",         "Groq API (LLaMa 3.1 8b / 3.3 70b)",       "Ultra-fast LLM inference"],
        ["Risk ML Model",        "XGBoost 3.2 + scikit-learn 1.8",          "Indian Loan Risk Predictor"],
        ["Vector DB – Users",    "ChromaDB 1.1",                            "Persistent user account & profile storage"],
        ["Vector DB – Policies", "FAISS 1.13 + sentence-transformers 5.2",  "Semantic policy document search (RAG)"],
        ["Web Scraping",         "duckduckgo-search (DDGS)",                "Boot-time live interest rate fetching"],
        ["Authentication",       "PyJWT + Bcrypt + passlib",                "JWT-based secure auth"],
        ["Data Processing",      "pandas 2.3 + numpy 2.4",                  "Feature engineering for ML model"],
        ["Logging",              "Custom structured logger",                 "Structured JSON logs with event codes"],
    ],
    col_widths=[1.8, 2.2, 2.8]
)

doc.add_paragraph()
heading("2.2  Frontend", level=2)
table(
    ["Layer", "Technology", "Purpose"],
    [
        ["Framework",       "React 18 + Vite",         "Component-based SPA with hot reload"],
        ["Routing",         "React Router v6",          "Client-side navigation"],
        ["Real-time",       "Native WebSocket API",     "Zero-latency bidirectional chat"],
        ["Styling",         "Vanilla CSS",              "Custom glassmorphism dark design system"],
        ["State",           "React Context API",        "Global JWT auth state"],
        ["Markdown",        "react-markdown",           "Render agent Markdown responses as rich UI"],
    ],
    col_widths=[1.5, 2.2, 3.1]
)


# ═════════════════════════════════════════════════════════════════════════════
# 3. SYSTEM ARCHITECTURE
# ═════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading("3. System Architecture")
para(
    "The system follows a layered architecture where the React frontend communicates with the FastAPI backend "
    "via REST for CRUD operations and a WebSocket for real-time chat. Every chat message flows through a "
    "LangGraph orchestrator which routes it to one or more CrewAI agents, each of which calls Python tool "
    "functions wrapping ML models and RAG pipelines."
)
doc.add_paragraph()

heading("3.1  Architecture Layers", level=2)
code_block([
    "  React Frontend (localhost:5173)",
    "        │  REST HTTP  /  WebSocket",
    "        ▼",
    "  FastAPI Backend (localhost:8000)",
    "     ├── Auth Router         (/auth/*)",
    "     ├── Profile Router      (/profile/*)",
    "     ├── Chat Router         (/chat/* + /ws/chat/*)",
    "     ├── Document Router     (/documents/*)",
    "     └── Guest Router        (/guest/*)",
    "              │",
    "              ▼",
    "  LangGraph Orchestrator  ──────────────────────────────────────",
    "    [intent_classifier] → [profile_collector] → [credit_analysis]",
    "    → [loan_matching] → [compliance_check] → [response_formatter]",
    "    └─ [rag_search] → [response_formatter]",
    "    └─ [improvement_plan] → [response_formatter]",
    "              │",
    "    ┌─────────┼──────────┐",
    "    ▼         ▼          ▼",
    "  XGBoost   FAISS      ChromaDB",
    "  (Risk)   (Policies)  (Users)",
])

doc.add_paragraph()
heading("3.2  Data Flow Summary", level=2)
para(
    "Every user message travels through the following pipeline before a response is returned:"
)
steps = [
    ("WebSocket Receive", "chat_router.py receives the raw JSON message with user content, JWT token, and session ID."),
    ("JWT Validation", "jwt_handler.py validates the bearer token before any processing begins."),
    ("Profile Fetch", "auth/service.py retrieves the authenticated user's full financial profile from ChromaDB."),
    ("Graph Invocation", "LangGraph.invoke() is called with an initial OrchestratorState dictionary."),
    ("Intent Classification", "The first node calls Groq LLaMa-3 to classify the user's intent."),
    ("Conditional Routing", "The graph routes to the appropriate node branch based on the classified intent."),
    ("Agent Execution", "One or more CrewAI agents are spawned, each calling Python tools for math and search."),
    ("Response Formatting", "The response_formatter node assembles all agent outputs into Markdown."),
    ("WebSocket Send", "The formatted Markdown string is sent back to the browser over the WebSocket connection."),
]
for i, (step, desc) in enumerate(steps, 1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    r = p.add_run(f"Step {i}: {step}  — ")
    set_font(r, bold=True, size=10.5, color=(31, 73, 125))
    r2 = p.add_run(desc)
    set_font(r2, size=10.5)


# ═════════════════════════════════════════════════════════════════════════════
# 4. SERVER BOOT SEQUENCE
# ═════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading("4. Server Boot Sequence")
para(
    "When uvicorn starts, the FastAPI lifespan context manager executes four sequential initialization "
    "steps before accepting any connections. This boot-time loading ensures zero latency during active chat sessions."
)
doc.add_paragraph()

heading("4.1  Step 1 — Live Rate Scatter-Gather Engine", level=2)
para("File: backend/ml_models/rate_predictor.py")
para(
    "The system autonomously scrapes real-world interest rates from the internet. Instead of querying a single "
    "search, it performs 3 targeted web searches (called 'chunks') to cover all loan categories and providers:"
)
table(
    ["Chunk", "Search Query", "Target Banks / NBFCs"],
    [
        ["A – Retail Loans",     "home/personal/car loan rates 2026 India",      "SBI, HDFC, ICICI, Bajaj Finserv, Axis"],
        ["B – Specialized",      "business/education/MSME loan rates 2026",       "Kotak, Bank of Baroda, Avanse, BoI"],
        ["C – Asset & Micro",    "gold loan / MUDRA rates Muthoot Manappuram",    "Muthoot Finance, Manappuram, IIFL"],
    ],
    col_widths=[1.5, 2.8, 2.5]
)
doc.add_paragraph()
para(
    "For each chunk: raw web text is extracted → fed to Groq LLaMa-3 with a strict JSON extraction prompt → "
    "the LLM returns a structured dictionary of bank_name→rate pairs. All 3 results are merged and used to "
    "recalculate the spread over the current RBI Repo Rate, dynamically overwriting the RATE_BANDS cache."
)
para("A time.sleep(2) delay between each chunk prevents IP-based rate limiting from the search API.", bold=False)

heading("4.2  Step 2 — XGBoost Risk Model", level=2)
para("File: backend/ml_models/fraud_detector.py")
para(
    "The pre-trained XGBoost model (loan_risk_model.joblib) is loaded from disk. This model was trained on "
    "5,000+ synthetic Indian borrower profiles and predicts the probability of loan default given 10 financial "
    "features. Once loaded, it stays in memory for instant inference during every chat session."
)

heading("4.3  Step 3 — RAG Pipeline (FAISS)", level=2)
para("File: backend/rag/pipeline.py")
para(
    "The FAISS vector index for policy search is loaded from disk. If no index exists, the pipeline reads "
    "all text files in backend/data/ (RBI guidelines, PMAY criteria, MUDRA scheme details, document checklists), "
    "creates sentence embeddings using all-MiniLM-L6-v2 from sentence-transformers, and builds a FAISS Flat-L2 "
    "index. This index supports semantic similarity search for user policy questions."
)

heading("4.4  Step 4 — LangGraph Compilation", level=2)
para("File: backend/orchestrator/graph.py")
para(
    "The LangGraph StateGraph is compiled with all 8 nodes and their conditional edges. This compilation validates "
    "the entire graph topology and caches the execution plan. Any misconfiguration would cause a hard failure at "
    "this point, making the system fail-fast before serving any users."
)


# ═════════════════════════════════════════════════════════════════════════════
# 5. LANGGRAPH ORCHESTRATOR — DETAILED FLOW
# ═════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading("5. LangGraph Orchestrator — Node-by-Node Flow")
para(
    "The orchestrator is a compiled LangGraph StateGraph. Every message enters at the intent_classifier node "
    "and follows conditional edges to subsequent nodes based on the values written to the shared state dictionary."
)

heading("5.1  OrchestratorState Schema", level=2)
table(
    ["Field", "Type", "Description"],
    [
        ["messages",            "List[BaseMessage]",  "Full LangChain message history"],
        ["intent",              "str",                "Classified intent: loan_inquiry / policy_question / general"],
        ["loan_request",        "dict",               "Extracted loan_amount, loan_type, tenure_months"],
        ["user_profile",        "dict",               "User's financial profile from ChromaDB"],
        ["credit_profile",      "dict",               "Output of credit analysis (DTI, risk tier, eligible)"],
        ["loan_products",       "list",               "Compared loan products from Loan Advisor agent"],
        ["compliance_result",   "dict",               "Eligibility, PMAY/MUDRA flags, required documents"],
        ["improvement_plan",    "dict",               "Score improvement steps for ineligible users"],
        ["agent_response",      "str",                "Final formatted Markdown response for the user"],
        ["flow",                "str",                "Internal flow state for routing decisions"],
        ["conversation_history","list",               "Serialized chat history for LLM context"],
    ],
    col_widths=[1.8, 1.6, 3.4]
)

doc.add_paragraph()
heading("5.2  Intent Classification Node", level=2)
para("File: orchestrator/intent_classifier.py  |  Node: intent_classifier")
para(
    "Uses Groq LLaMa-3.1-8b-instant (smallest, fastest model) to parse the user's natural language message "
    "and return a strict JSON object. The prompt enforces four intent categories:"
)
table(
    ["Intent", "Trigger Condition", "Graph Route"],
    [
        ["loan_inquiry",     "User explicitly requests a loan amount/type for a new calculation",  "→ profile_collector"],
        ["policy_question",  "User asks about government schemes, RBI rules, document requirements", "→ rag_search"],
        ["profile_update",   "User provides personal financial data to update their profile",       "→ response_formatter"],
        ["general",          "Greetings, arguments ('but you said...'), or off-topic questions",    "→ response_formatter (LLM chat)"],
    ],
    col_widths=[1.5, 2.8, 2.5]
)
doc.add_paragraph()
para(
    "The classifier also extracts loan_amount (with Lakh/Crore conversion), loan_type (from keywords like "
    "'house'→home_loan, 'boutique'→business_loan), and tenure_months from the message. It applies a "
    "sanity cap: any amount > ₹100 Crore is rejected as a likely LLM hallucination."
)
para(
    "Fallback: if Groq is unavailable, a regex heuristic engine runs locally to classify intent "
    "without any external API call, ensuring 100% uptime.",
    bold=False
)

heading("5.3  Profile Collector Node", level=2)
para("File: orchestrator/nodes.py  |  Node: profile_collector")
para(
    "Checks whether the authenticated user's profile has all required financial fields: annual_income, "
    "credit_score, employment_months, existing_loans, existing_emi_amount, credit_utilization, city, and age. "
    "If any are missing, it routes to response_formatter with a request for the user to complete their profile "
    "via the Profile tab. If complete, it routes to credit_analysis."
)

heading("5.4  Credit Analysis Node", level=2)
para("File: agents/credit_analyst.py  |  Node: credit_analysis  |  Model: llama-3.1-8b-instant")
para("This node spawns a CrewAI Credit Analyst agent with one tool: loan_risk_scorer. The workflow:")
step_list = [
    "Agent calls loan_risk_scorer(annual_income, credit_score, age, existing_emi, loan_amount, loan_type)",
    "XGBoost model runs inference → returns risk_score (0–1) and risk_tier (Low / Medium / High)",
    "DETERMINISTIC OVERRIDE: After CrewAI completes, Python recalculates DTI = (existing_emi / (annual_income/12)) × 100",
    "Eligibility check: CIBIL ≥ 600 AND DTI ≤ 50% → eligible=True, else eligible=False",
    "If eligible → loan_matching node. If ineligible → improvement_plan node.",
]
for s in step_list:
    bullet(s)

heading("5.5  Loan Matching Node", level=2)
para("File: agents/loan_advisor.py  |  Node: loan_matching  |  Model: llama-3.3-70b-versatile")
para("The Loan Advisor agent has two tools: rate_predictor and emi_calculator.")
para("The agent flow:")
step_list2 = [
    "Calls rate_predictor(loan_type, credit_score) → reads from RAM cache → returns {min_rate, max_rate, personalized_rate, provider_rates}",
    "For each provider in provider_rates, calls emi_calculator(principal, rate, tenure_months) → Python math → returns {monthly_emi, total_interest, total_repayment}",
    "Constructs a comparison table of 3–4 bank options with different rates and tenures",
    "If CrewAI/Groq fails → _deterministic_advisory() fallback runs the exact same logic in pure Python, generating a complete table without any LLM",
]
for s in step_list2:
    bullet(s)

heading("5.6  Compliance Check Node", level=2)
para("File: agents/risk_compliance.py  |  Node: compliance_check  |  Model: llama-3.3-70b-versatile")
para("The Risk Compliance agent has one tool: policy_search (FAISS RAG). The agent:")
step_list3 = [
    "Calls policy_search('RBI policies for [loan_type]') → FAISS semantic search → top 3 chunks returned",
    "Calls policy_search('PMAY eligibility criteria') → FAISS search",
    "Calls policy_search('MUDRA scheme eligibility for [loan_type]') → FAISS search",
    "Synthesizes all retrieved chunks to determine: eligibility (yes/no), approval_likelihood (%), PMAY/MUDRA applicability, required documents list",
]
for s in step_list3:
    bullet(s)

heading("5.7  Response Formatter Node", level=2)
para("File: orchestrator/nodes.py  |  Node: response_formatter")
para(
    "The final node in every flow. It assembles all agent outputs into a single, richly formatted Markdown "
    "string that the React frontend renders as a structured UI card. For loan_inquiry responses, it generates "
    "a fixed-format card with sections: Credit Analysis, Loan Options Comparison table, Recommendation, "
    "Compliance & Eligibility, and Required Documents."
)
para(
    "For general/greeting intents, it calls Groq LLaMa-3.1-8b-instant with the user's chat history and profile "
    "name to generate a natural conversational response — without running any financial tools.",
    bold=False
)


# ═════════════════════════════════════════════════════════════════════════════
# 6. ML MODELS
# ═════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading("6. Machine Learning Models")

heading("6.1  XGBoost Indian Loan Risk Model", level=2)
para("Files: ml_models/fraud_detector.py, ml_models/build_loan_risk_model.py")
para("The model classifies whether a loan application is High Risk, Medium Risk, or Low Risk.")
table(
    ["Feature", "Type", "Description"],
    [
        ["credit_score",          "int",   "CIBIL credit score (300–900)"],
        ["annual_income",         "float", "Annual income in INR"],
        ["loan_amount",           "float", "Requested loan amount in INR"],
        ["loan_to_income_ratio",  "float", "loan_amount / annual_income"],
        ["employment_months",     "int",   "Months of employment history"],
        ["existing_emi_amount",   "float", "Current monthly EMI obligations"],
        ["debt_to_income_ratio",  "float", "existing_emi × 12 / annual_income"],
        ["credit_utilization",    "float", "Credit card utilization percentage"],
        ["age",                   "int",   "Borrower age in years"],
        ["loan_type_encoded",     "int",   "Integer-encoded loan type"],
    ],
    col_widths=[2.0, 0.8, 4.0]
)
doc.add_paragraph()
table(
    ["Attribute", "Value"],
    [
        ["Algorithm",             "XGBoost Classifier"],
        ["Training Data",         "5,000+ synthetic Indian borrower profiles"],
        ["Output",                "risk_score (0–1) and risk_tier (Low / Medium / High)"],
        ["Saved Model",           "ml_models/saved/loan_risk_model.joblib"],
        ["Load Time",             "< 0.5 seconds from disk"],
        ["Inference Time",        "< 10 milliseconds per prediction"],
    ],
    col_widths=[2.2, 4.6]
)

heading("6.2  EMI Calculator", level=2)
para("File: ml_models/emi_calculator.py")
para(
    "Pure Python implementation of the standard Indian EMI formula. This function is called by the CrewAI "
    "Loan Advisor agent for every bank product in the comparison table. The LLM never computes EMI values — "
    "it only passes the parameters as tool arguments."
)
code_block([
    "  EMI = P × R × (1 + R)^N",
    "        ─────────────────",
    "          (1 + R)^N − 1",
    "",
    "  Where:",
    "    P = Principal loan amount (INR)",
    "    R = Monthly interest rate (annual_rate / 12 / 100)",
    "    N = Tenure in months",
    "",
    "  Returns: monthly_emi, total_interest, total_repayment",
])

heading("6.3  Rate Predictor & RATE_BANDS Cache", level=2)
para("File: ml_models/rate_predictor.py")
para(
    "At boot time, the Scatter-Gather engine updates a global in-memory dictionary called RATE_BANDS. "
    "This dictionary maps each loan type to its current spread (in percentage points above the RBI Repo Rate) "
    "and a list of provider names. The predict_rate() function performs the following at query time:"
)
step_list4 = [
    "1. Look up RATE_BANDS[loan_type] → get {spread, providers}",
    "2. Apply credit score tier modifier: Excellent (750+) = −0.05%, Poor (<600) = +2.0%",
    "3. Compute personalized_rate = repo_rate + spread + tier_modifier",
    "4. Return {min_rate, max_rate, personalized_rate, provider_rates} to the calling agent tool",
]
for s in step_list4:
    para(s, indent=True)


# ═════════════════════════════════════════════════════════════════════════════
# 7. RAG PIPELINE
# ═════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading("7. Retrieval-Augmented Generation (RAG) Pipeline")
para("File: rag/pipeline.py")
para(
    "The RAG pipeline enables the system to answer policy questions grounded in actual documents, "
    "preventing the AI from inventing government scheme rules or eligibility criteria."
)

heading("7.1  Indexing Phase (Boot-Time or First Run)", level=2)
step_list5 = [
    "Source documents are plain text files placed in backend/data/",
    "LangChain RecursiveCharacterTextSplitter splits documents into 500-token chunks with 50-token overlap",
    "All chunks are embedded using sentence-transformers all-MiniLM-L6-v2 (384-dimensional vectors)",
    "Embeddings are stored in a FAISS Flat-L2 index and persisted to disk at rag/faiss_index/",
    "On subsequent boots, the saved index is loaded from disk (skipping re-ingestion)",
]
for s in step_list5:
    bullet(s)
doc.add_paragraph()
para("Policy documents indexed:")
docs = [
    "RBI lending guidelines and credit score classification rules",
    "PMAY (Pradhan Mantri Awas Yojana) — eligibility, income limits, subsidy rates",
    "MUDRA / PMEGP scheme details — Shishu / Kishore / Tarun categories and loan limits",
    "Standard loan documentation requirements per loan type",
]
for d in docs:
    bullet(d)

heading("7.2  Query Phase (Per User Request)", level=2)
step_list6 = [
    "User query is embedded with the same all-MiniLM-L6-v2 model",
    "FAISS performs cosine similarity search → top 5 most relevant chunks retrieved",
    "Retrieved chunks are formatted with source attribution and relevance score",
    "Chunks + user query are sent to Groq LLaMa-3.3-70b-versatile for synthesis",
    "LLM generates a coherent, contextual answer grounded only in the retrieved chunks",
    "If LLM fails → raw chunks are formatted and returned directly as a fallback",
]
for s in step_list6:
    bullet(s)


# ═════════════════════════════════════════════════════════════════════════════
# 8. AUTHENTICATION
# ═════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading("8. Authentication & User Management")
para("Files: auth/service.py, auth/jwt_handler.py, routers/auth_router.py")

heading("8.1  User Storage (ChromaDB)", level=2)
para(
    "All user accounts and financial profiles are stored in ChromaDB — a vector database. "
    "Each user document contains:"
)
table(
    ["Field", "Type", "Description"],
    [
        ["user_id",                 "UUID str", "Auto-generated unique identifier"],
        ["username",                "str",      "Login username"],
        ["full_name",               "str",      "Display name"],
        ["hashed_password",         "str",      "Bcrypt-hashed password"],
        ["annual_income",           "float",    "Annual income in INR"],
        ["credit_score",            "int",      "CIBIL score 300–900"],
        ["employment_months",       "int",      "Months of work experience"],
        ["existing_loans",          "int",      "Number of active loans"],
        ["existing_emi_amount",     "float",    "Monthly EMI obligations in INR"],
        ["credit_utilization",      "float",    "Credit card utilization %"],
        ["age",                     "int",      "Age in years"],
        ["city",                    "str",      "City of residence"],
        ["loan_type_interest",      "str",      "Preferred loan type"],
        ["profile_complete",        "bool",     "Whether all required fields are filled"],
    ],
    col_widths=[2.0, 1.0, 3.8]
)

heading("8.2  JWT Authentication Flow", level=2)
step_list7 = [
    "User POSTs username + password to /auth/login",
    "Service fetches user from ChromaDB by username",
    "passlib/bcrypt verifies the plaintext password against the stored hash",
    "On success: PyJWT creates a signed token with {sub: user_id, exp: 24 hours}",
    "Frontend stores the token in React Context (in-memory, not localStorage)",
    "Every subsequent request includes Authorization: Bearer <token> header",
    "FastAPI dependency get_current_user() validates the token on every protected route",
]
for s in step_list7:
    bullet(s)


# ═════════════════════════════════════════════════════════════════════════════
# 9. FRONTEND
# ═════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading("9. Frontend Architecture")

heading("9.1  Pages", level=2)
table(
    ["Page", "File", "Description"],
    [
        ["Login / Signup",  "LoginPage.jsx",   "JWT-authenticated login and registration forms"],
        ["Dashboard",       "DashboardPage.jsx","Full application shell: sidebar, chat panel, profile editor, tools"],
        ["Guest",           "GuestPage.jsx",    "Demo experience with sample financial profile, no login required"],
    ],
    col_widths=[1.5, 2.0, 3.3]
)

heading("9.2  Components", level=2)
table(
    ["Component", "File", "Description"],
    [
        ["ChatPanel",       "ChatPanel.jsx",       "WebSocket lifecycle, message send/receive, react-markdown rendering"],
        ["CreditGauge",     "CreditGauge.jsx",     "Animated SVG circular gauge showing credit score 300–900"],
        ["EMICalculator",   "EMICalculator.jsx",   "Interactive standalone EMI calculator (calls backend API)"],
        ["DocumentUpload",  "DocumentUpload.jsx",  "Drag-and-drop document uploader that auto-extracts profile fields"],
    ],
    col_widths=[1.8, 2.0, 3.0]
)

heading("9.3  WebSocket Chat Protocol", level=2)
para("The frontend opens a persistent WebSocket at: ws://localhost:8000/ws/chat/{user_id}")
para("Each message sent to the server:")
code_block([
    "  {",
    '    "message":    "I want a home loan of 50 lakhs",',
    '    "token":      "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9...",',
    '    "session_id": "1e64b76f-574e-4d16-a490-40c98ea7be9d"',
    "  }",
])
doc.add_paragraph()
para("The server validates the token, runs the LangGraph pipeline, and sends back the final Markdown string.")
para(
    "On disconnection, the frontend automatically reconnects with exponential backoff, ensuring a seamless "
    "experience even if the backend briefly restarts.",
    bold=False
)

heading("9.4  Chat Session Management", level=2)
step_list8 = [
    "Sessions are listed in the sidebar with title, relative timestamp, and last message preview",
    "Clicking '+ New' creates a new session via POST /chat/sessions/{user_id}",
    "Switching sessions fetches the full message history via GET /chat/sessions/{user_id}/{session_id}/messages",
    "Clicking the ✕ button on a session opens a glassmorphism modal with blurred background for delete confirmation",
    "Confirming delete calls DELETE /chat/sessions/{user_id}/{session_id} → removes from backend JSON store permanently",
    "Sessions persist across browser sessions; history is reloaded on login",
]
for s in step_list8:
    bullet(s)

heading("9.5  Design System", level=2)
para(
    "The frontend uses a custom glassmorphism dark-mode design system defined in index.css. "
    "Key design tokens:"
)
table(
    ["Token", "Value", "Usage"],
    [
        ["--primary",       "#6366f1 (Indigo)",    "Buttons, active states, highlights"],
        ["--surface",       "rgba(15,15,35,0.95)",  "Main background panels"],
        ["--surface-2",     "rgba(20,20,45,0.98)",  "Cards, modals"],
        ["--danger",        "#dc2626 (Red)",         "Delete buttons, error states"],
        ["--success",       "#22c55e (Green)",       "Eligible indicators, success messages"],
        ["--border",        "rgba(255,255,255,0.1)", "Subtle separators"],
        ["--text-muted",    "rgba(255,255,255,0.5)", "Secondary text, timestamps"],
    ],
    col_widths=[1.8, 2.2, 2.8]
)


# ═════════════════════════════════════════════════════════════════════════════
# 10. ENTERPRISE GUARDRAILS
# ═════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading("10. Enterprise Guardrails & Safety Systems")
para(
    "The system implements 7 distinct safety mechanisms to ensure accuracy, reliability, "
    "and regulatory safety in a financial advisory context."
)

guardrails = [
    (
        "1. Mathematical Determinism (Anti-Hallucination)",
        "The LLM is structurally prevented from performing any financial calculations. "
        "It can only invoke Python tool functions and format their outputs. "
        "Every EMI, DTI ratio, risk score, and interest rate is computed by deterministic Python code.",
        "crewai_tools.py, emi_calculator.py, rate_predictor.py"
    ),
    (
        "2. Intent Domain-Lock (Anti-Jailbreak)",
        "The Intent Classifier uses a strict allowlist of 4 intent categories. "
        "Anything not matching loan_inquiry, policy_question, or profile_update is routed to the lightweight "
        "general conversation handler, completely bypassing the AI financial agent pipeline.",
        "orchestrator/intent_classifier.py"
    ),
    (
        "3. Conversational Disambiguation",
        "The Intent Classifier system prompt explicitly instructs the LLM: 'Use general for arguments "
        "(but you said...) or clarifications that do not ask to rerun a new loan calculation.' "
        "This prevents follow-up questions from re-triggering the full 5-node agent pipeline.",
        "orchestrator/intent_classifier.py (SYSTEM_PROMPT)"
    ),
    (
        "4. Deterministic Fallback System",
        "Every CrewAI agent call is wrapped in try/except. If Groq API returns a 400 error, "
        "rate limit, or malformed JSON, _deterministic_advisory() is called instead. "
        "This function generates a complete loan comparison table using only Python math — "
        "the user never sees an error message.",
        "agents/loan_advisor.py, agents/credit_analyst.py"
    ),
    (
        "5. RAG-Grounded Compliance",
        "The Risk Compliance agent can only cite information from the FAISS vector index, "
        "which contains only administrator-uploaded policy documents. "
        "It cannot draw on the LLM's pre-trained memory for policy facts, "
        "preventing fabrication of government scheme rules.",
        "agents/risk_compliance.py, rag/pipeline.py"
    ),
    (
        "6. Boot-Time Rate Locking",
        "Interest rates are scraped once at server boot and locked in memory (RATE_BANDS dict). "
        "This prevents the system from using stale pre-trained LLM rates from years ago, "
        "while also avoiding the latency of live-searching during every user conversation. "
        "The system also applies time.sleep(2) between search chunks to prevent IP bans.",
        "ml_models/rate_predictor.py"
    ),
    (
        "7. Pydantic Schema Coercion",
        "All CrewAI tool functions use Pydantic validation with explicit type coercion. "
        "The credit_score parameter is defined as int | str with an explicit int() cast, "
        "preventing a litellm.BadRequestError crash when the LLM passes a string '800' instead of integer 800.",
        "tools/crewai_tools.py"
    ),
]

for title, desc, files in guardrails:
    p = doc.add_paragraph()
    r = p.add_run(title)
    set_font(r, bold=True, size=11, color=(31, 73, 125))
    para(desc, indent=True)
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Inches(0.3)
    r2 = p2.add_run(f"Implemented in: {files}")
    set_font(r2, size=9.5, italic=True, color=(100, 100, 100))
    doc.add_paragraph()


# ═════════════════════════════════════════════════════════════════════════════
# 11. END-TO-END EXAMPLE
# ═════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading("11. End-to-End Request Walkthrough")
para('User message: "I want a home loan of 50 lakhs"')
para("User profile: Pranav, annual_income=₹20L, credit_score=800, existing_emi=₹16,000/month")
doc.add_paragraph()

steps_e2e = [
    ("WS Receive",         "chat_router.py",                "WebSocket receives JSON: {message, token, session_id}"),
    ("JWT Validate",       "auth/jwt_handler.py",            "Token decoded → user_id = 111e5e6f-..."),
    ("Profile Fetch",      "auth/service.py",                "ChromaDB query → full financial profile returned"),
    ("Graph Start",        "orchestrator/graph.py",          "LangGraph.invoke({messages, user_profile, ...})"),
    ("Intent Classify",    "orchestrator/intent_classifier.py","Groq LLaMa-3.1-8b → {intent: loan_inquiry, amount: 5000000, type: home_loan}"),
    ("Profile Check",      "orchestrator/nodes.py",          "All 8 required profile fields present → eligible to proceed"),
    ("Credit Analysis",    "agents/credit_analyst.py",       "CrewAI agent → calls loan_risk_scorer → XGBoost → risk_score=0.26, tier=Low"),
    ("DTI Override",       "agents/credit_analyst.py",       "Python: DTI = (16000 / (2000000/12)) × 100 = 9.6% → eligible=True"),
    ("Rate Prediction",    "tools/crewai_tools.py",          "rate_predictor('home_loan', 800) → RAM cache → {min:7.6%, max:10.75%, personal:7.65%}"),
    ("EMI #1 (SBI)",       "ml_models/emi_calculator.py",    "calculate_emi(5000000, 7.6, 240) → ₹40,586/month, ₹4.74L interest"),
    ("EMI #2 (HDFC)",      "ml_models/emi_calculator.py",    "calculate_emi(5000000, 8.4, 240) → ₹43,075/month, ₹5.34L interest"),
    ("EMI #3 (ICICI)",     "ml_models/emi_calculator.py",    "calculate_emi(5000000, 7.65, 180) → ₹46,778/month, ₹3.42L interest"),
    ("Policy Search",      "rag/pipeline.py",                "FAISS search 'RBI home loan', 'PMAY eligibility' → top 3 chunks each"),
    ("Compliance",         "agents/risk_compliance.py",      "Groq synthesizes chunks → eligibility=YES, PMAY=YES, approval=90%"),
    ("Format Response",    "orchestrator/nodes.py",          "Markdown card assembled: Credit Analysis + Table + Compliance sections"),
    ("WS Send",            "routers/chat_router.py",         "Markdown string sent over WebSocket → browser renders formatted card"),
]

t = doc.add_table(rows=1, cols=4)
t.style = "Table Grid"
hdr = t.rows[0].cells
for i, h in enumerate(["Step", "Component", "File", "What Happens"]):
    hdr[i].text = h
    for run in hdr[i].paragraphs[0].runs:
        run.font.bold = True
        run.font.name = "Calibri"
        run.font.size = Pt(9.5)
    tc = hdr[i]._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "1F497D")
    tcPr.append(shd)
    for par in hdr[i].paragraphs:
        for run in par.runs:
            run.font.color.rgb = RGBColor(255, 255, 255)

for i, (comp, f, what) in enumerate(steps_e2e, 1):
    row = t.add_row().cells
    row[0].text = str(i)
    row[1].text = comp
    row[2].text = f
    row[3].text = what
    for cell in row:
        for run in cell.paragraphs[0].runs:
            run.font.name = "Calibri"
            run.font.size = Pt(9)
    if i % 2 == 0:
        for cell in row:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"),   "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"),  "EEF2FF")
            tcPr.append(shd)

for row in t.rows:
    row.cells[0].width = Inches(0.4)
    row.cells[1].width = Inches(1.4)
    row.cells[2].width = Inches(1.9)
    row.cells[3].width = Inches(3.1)

doc.add_paragraph()
para("Total response time: approximately 5–8 seconds (dominated by 3 sequential Groq LLM API calls).")


# ═════════════════════════════════════════════════════════════════════════════
# 12. API REFERENCE
# ═════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading("12. API Reference")

heading("12.1  REST Endpoints", level=2)
table(
    ["Method", "Endpoint", "Auth", "Description"],
    [
        ["POST", "/auth/signup",                                   "None",  "Create new user account"],
        ["POST", "/auth/login",                                    "None",  "Login → returns JWT token"],
        ["GET",  "/profile/{user_id}",                             "JWT",   "Retrieve user financial profile"],
        ["POST", "/profile/update",                                "JWT",   "Update profile fields"],
        ["GET",  "/chat/sessions/{user_id}",                       "JWT",   "List all chat sessions"],
        ["POST", "/chat/sessions/{user_id}",                       "JWT",   "Create new chat session"],
        ["DELETE","/chat/sessions/{user_id}/{session_id}",          "JWT",   "Permanently delete a session"],
        ["GET",  "/chat/sessions/{user_id}/{session_id}/messages",  "JWT",   "Get message history for session"],
        ["WS",   "/ws/chat/{user_id}",                             "JWT*",  "Real-time chat WebSocket"],
        ["POST", "/documents/upload/{user_id}",                    "JWT",   "Upload and parse financial documents"],
        ["GET",  "/health",                                        "None",  "System health check"],
        ["POST", "/logs",                                          "None",  "Frontend structured log ingestion"],
    ],
    col_widths=[0.7, 3.0, 0.7, 2.4]
)
para("* JWT is passed in the WebSocket message payload, not in the HTTP upgrade header.", indent=True)

heading("12.2  WebSocket Message Format", level=2)
para("Sent by client:")
code_block([
    '  { "message": "I want a home loan", "token": "eyJ...", "session_id": "uuid..." }',
])
para("Received from server:")
code_block([
    '  "## 🏦 Home Loan Assessment — ₹5,000,000\\n\\n### 📊 Credit Analysis..."',
    "  (plain Markdown string, rendered by react-markdown in the frontend)",
])


# ═════════════════════════════════════════════════════════════════════════════
# 13. FILE STRUCTURE
# ═════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading("13. Project File Structure")
code_block([
    "loanv2/",
    "├── README.md                         Simple setup guide",
    "├── Project_Documentation.docx        This document",
    "│",
    "├── backend/",
    "│   ├── .env                          Environment variables (never commit to Git)",
    "│   ├── main.py                       FastAPI app + boot lifespan + middleware",
    "│   ├── config.py                     Pydantic settings loader from .env",
    "│   ├── logger.py                     Structured JSON logger with event codes",
    "│   ├── requirements.txt              All Python dependencies",
    "│   │",
    "│   ├── auth/",
    "│   │   ├── service.py                ChromaDB user store + bcrypt auth + profile CRUD",
    "│   │   └── jwt_handler.py            JWT creation and validation",
    "│   │",
    "│   ├── agents/",
    "│   │   ├── credit_analyst.py         CrewAI agent: risk scoring + DTI override",
    "│   │   ├── loan_advisor.py           CrewAI agent: rate prediction + EMI calc + fallback",
    "│   │   └── risk_compliance.py        CrewAI agent: RAG policy search + eligibility",
    "│   │",
    "│   ├── orchestrator/",
    "│   │   ├── graph.py                  LangGraph state machine (nodes + conditional edges)",
    "│   │   ├── nodes.py                  All 8 node execution functions",
    "│   │   ├── intent_classifier.py      Groq-powered intent classifier with regex fallback",
    "│   │   └── state.py                  OrchestratorState TypedDict definition",
    "│   │",
    "│   ├── ml_models/",
    "│   │   ├── rate_predictor.py         Scatter-Gather web scraper + RATE_BANDS cache",
    "│   │   ├── emi_calculator.py         Pure Python EMI math (P×R×(1+R)^N / ((1+R)^N-1))",
    "│   │   ├── fraud_detector.py         XGBoost Loan Risk Predictor wrapper",
    "│   │   └── build_loan_risk_model.py  Model training script (run once to generate .joblib)",
    "│   │",
    "│   ├── tools/",
    "│   │   └── crewai_tools.py           @crewai_tool wrappers: loan_risk_scorer, emi_calculator,",
    "│   │                                  rate_predictor, policy_search",
    "│   │",
    "│   ├── rag/",
    "│   │   └── pipeline.py               FAISS index build/load + semantic search + LLM synthesis",
    "│   │",
    "│   ├── routers/",
    "│   │   ├── auth_router.py            POST /auth/signup, POST /auth/login",
    "│   │   ├── profile_router.py         GET/POST /profile/*",
    "│   │   ├── chat_router.py            GET/POST/DELETE /chat/* + WS /ws/chat/*",
    "│   │   ├── document_router.py        POST /documents/upload/*",
    "│   │   ├── guest_router.py           Guest demo mode endpoints",
    "│   │   └── logs_router.py            POST /logs (frontend log ingestion)",
    "│   │",
    "│   ├── data/                         Policy text files for RAG ingestion",
    "│   ├── chromadb_storage/             ChromaDB persistent user data (auto-created)",
    "│   └── uploads/                      Uploaded financial documents (auto-created)",
    "│",
    "└── frontend/",
    "    ├── package.json",
    "    ├── vite.config.js",
    "    └── src/",
    "        ├── main.jsx                  React entry point",
    "        ├── App.jsx                   React Router setup",
    "        ├── index.css                 4,000+ line glassmorphism CSS design system",
    "        ├── context/",
    "        │   └── AuthContext.jsx       Global JWT auth state (login, logout, user object)",
    "        ├── api/",
    "        │   └── client.js             apiGet, apiPost, sendLog helper functions",
    "        ├── pages/",
    "        │   ├── DashboardPage.jsx     Main shell + sidebar + chat session management + modal",
    "        │   ├── LoginPage.jsx         Auth forms (login + signup tab)",
    "        │   └── GuestPage.jsx         Demo mode with sample profile",
    "        └── components/",
    "            ├── ChatPanel.jsx         WebSocket client + message list + Markdown rendering",
    "            ├── CreditGauge.jsx       Animated SVG circular credit score gauge",
    "            ├── EMICalculator.jsx     Interactive standalone EMI tool",
    "            └── DocumentUpload.jsx    Drag-and-drop document parser",
])


# ═════════════════════════════════════════════════════════════════════════════
# SAVE
# ═════════════════════════════════════════════════════════════════════════════
output_path = r"C:\Users\TammaSatyaPranavMAQS\OneDrive - MAQ Software\Documents\loanv2\Project_Documentation.docx"
doc.save(output_path)
print(f"✅  Document saved → {output_path}")
